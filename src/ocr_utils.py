import concurrent
import os
import easyocr
import pytesseract
from PIL import ImageEnhance, ImageFilter
from concurrent.futures import ThreadPoolExecutor, TimeoutError
from docx import Document
import docx2txt
from pdf2image import convert_from_path
import zipfile
import shutil
import hashlib
from src.image_azure_caption import process_image_caption  # Import the image captioning function

# Path to Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r"D:\Tesseract-OCR\tesseract.exe"

def generate_unique_name(file_path):
    return hashlib.md5(file_path.encode()).hexdigest()

def preprocess_image(image):
    """Preprocess an image (enhance contrast and sharpness, denoise)"""
    image = image.convert("RGB")
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(1.2)  # Adjust contrast enhancement
    return image

def process_image_tesseract(image):
    try:
        preprocessed_image = preprocess_image(image)
        text = pytesseract.image_to_string(preprocessed_image)
        return text
    except Exception as e:
        print(f"An error occurred during Tesseract OCR processing: {e}")
        return ""

def process_image_easyocr(image):
    try:
        preprocessed_image = preprocess_image(image)
        reader = easyocr.Reader(lang_list=['en', 'hi', 'mr'])
        results = reader.readtext(preprocessed_image)
        text = " ".join([result[1] for result in results])
        return text
    except Exception as e:
        print(f"An error occurred during EasyOCR processing: {e}")
        return ""

def extract_text_from_docx(docx_path):
    try:
        extracted_text = ""
        text = docx2txt.process(docx_path)
        extracted_text += text
        doc = Document(docx_path)
        for table in doc.tables:
            table_text = extract_table_text(table)
            extracted_text += table_text + "\n"
        return extracted_text
    except Exception as e:
        print(f"An error occurred during DOCX processing: {e}")
        return ""

def extract_table_text(table):
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            cell_text = ""
            for paragraph in cell.paragraphs:
                cell_text += paragraph.text + " "
            table_text += cell_text.strip() + "\t"
        table_text += "\n"
    return table_text

def process_pdf(pdf_path, poppler_path, output_folder):
    try:
        pdf_filename = os.path.basename(pdf_path)
        images = convert_from_path(pdf_path, poppler_path=poppler_path)
        def process_single_image(image, engine):
            try:
                if engine == 'tesseract':
                    text = process_image_tesseract(image)
                else:
                    text = process_image_easyocr(image)
                return text
            except TimeoutError:
                return ""
        extracted_text = ""
        current_engine = 'tesseract'
        engine_switch_time = 15  # Switch OCR engine after 15 seconds
        with ThreadPoolExecutor() as executor:
            for image in images:
                future = executor.submit(process_single_image, image, current_engine)
                try:
                    text = future.result(timeout=engine_switch_time)
                except TimeoutError:
                    current_engine = 'easyocr' if current_engine == 'tesseract' else 'tesseract'
                    text = future.result()
                extracted_text += text
        return extracted_text
    except Exception as e:
        print(f"An error occurred during PDF processing: {e}")
        return ""

def extract_images_from_docx(docx_file):
    temp_dir = "temp_extracted_images"
    os.makedirs(temp_dir, exist_ok=True)
    temp_zip_file = os.path.join(temp_dir, "temp_docx.zip")
    shutil.copy(docx_file, temp_zip_file)
    with zipfile.ZipFile(temp_zip_file, "r") as zip_file:
        image_dir = os.path.join(temp_dir, "images")
        os.makedirs(image_dir, exist_ok=True)
        for file_info in zip_file.infolist():
            if "word/media/" in file_info.filename:
                filename = os.path.basename(file_info.filename)
                image_data = zip_file.read(file_info.filename)
                with open(os.path.join(image_dir, filename), "wb") as image_file:
                    image_file.write(image_data)
    os.remove(temp_zip_file)
    return image_dir

def process_images_with_caption(image_dir):
    image_files = os.listdir(image_dir)
    image_captions = []
    with concurrent.futures.ThreadPoolExecutor() as executor:
        image_paths = [os.path.join(image_dir, image_file) for image_file in image_files]
        image_captions = list(executor.map(process_image_caption, image_paths))
    return image_captions